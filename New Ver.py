import sys
import json
import datetime
import logging
import traceback
import tkinter as tk
from tkinter import ttk, simpledialog, messagebox, filedialog
from pathlib import Path
import time
import queue
import threading
import re # Added for regular expressions to extract timestamp

import mss
import mss.tools

from docx import Document
from docx.shared import Inches
from screeninfo import get_monitors
import requests
import keyboard

# For displaying images fullscreen
from PIL import Image, ImageTk
import cv2 # pip install opencv-python
import numpy as np # Needed for PIL to OpenCV conversion

# ── IMPORTANT: CENTRALIZED BOT TOKEN ──────────────────────────────────────────
# REPLACE THIS WITH YOUR ACTUAL TELEGRAM BOT TOKEN.
# THIS TOKEN IS HARDCODED AND NOT SAVED IN config.json FOR SECURITY REASONS.
TELEGRAM_BOT_TOKEN = "8070792966:AAEyKF3g2SUNntmDJIctc8InyprORYOj4xg" 
# ──────────────────────────────────────────────────────────────────────────────

# ── Logging Setup ─────────────────────────────────────────────────────────────
log_dir = Path.cwd() / "logs"
log_dir.mkdir(parents=True, exist_ok=True)
log_file_path = log_dir / "app.log"

logging.basicConfig(
    filename=log_file_path,
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(message)s"
)
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO) 
console_formatter = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
console_handler.setFormatter(console_formatter)
logging.getLogger().addHandler(console_handler)

logging.info("===== Script start =====")

# ── Paths & Config ────────────────────────────────────────────────────────────
BASE_DIR = Path.home() / "Documents" / "Trading Journal"
CONFIG_FILE = "config.json"

SCRIPT_DIR = Path(sys.argv[0]).resolve().parent
# ── MODIFIED: Centralized Icon Path and Window Title ──────────────────────────
ICON_PATH = SCRIPT_DIR / "app_icon.ico" # Make sure you place your .ico file here and name it app_icon.ico
MAIN_WINDOW_TITLE = "Trading Journal Screenshot Tool" # Your desired title for the main GUI window
# ──────────────────────────────────────────────────────────────────────────────

telegram_queue = queue.Queue()

def get_base_path() -> Path:
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    return BASE_DIR

def get_cfg_path() -> Path:
    return SCRIPT_DIR / CONFIG_FILE

def load_cfg() -> dict:
    p = get_cfg_path()
    cfg = {}
    if p.exists():
        try:
            cfg = json.loads(p.read_text("utf-8"))
        except json.JSONDecodeError:
            logging.error(f"Error decoding config.json, creating a new one. Original content: {p.read_text('utf-8')}")
            messagebox.showwarning("Config Error", "Configuration file is corrupted. A new one will be created.")
    
    cfg.setdefault("telegram_chat_id", "")
    cfg.setdefault("instrument", "6B")
    cfg.setdefault("enable_telegram_send", True)
    cfg.setdefault("monitor_names", ",".join(f"Monitor {i+1}" for i in range(len(get_monitors()))))
    cfg.setdefault("default_description", "Reviewing trade setup.")
    cfg.setdefault("last_view_path", str(get_base_path())) # New config for last viewed path
    return cfg

def save_cfg(cfg: dict):
    p = get_cfg_path()
    p.write_text(json.dumps(cfg, ensure_ascii=False, indent=4), "utf-8")
    logging.info(f"Configuration saved to {p}")

# ── Globals & Constants ───────────────────────────────────────────────────────
monitors = get_monitors()
INSTRUMENTS = [
    "6E","6B","6A","6N","6S","6J","6C",
    "ES","NQ","YM","CL","GC","SI","ZB","ZN","ZF"
]

# ── Date Helpers ──────────────────────────────────────────────────────────────
def get_season(m: int) -> str:
    if m in (12,1,2): return "Winter"
    if m in (3,4,5):  return "Spring"
    if m in (6,7,8):  return "Summer"
    return "Autumn"

def get_week_of_month(d: datetime.date) -> int:
    first = d.replace(day=1)
    return d.isocalendar()[1] 

def get_save_directory(event: str, inst: str, now: datetime.datetime) -> Path:
    base = get_base_path()
    y = str(now.year)
    mname = now.strftime("%B")
    season = get_season(now.month)
    week = f"Week_{now.isocalendar()[1]}" 
    day = now.strftime("%Y-%m-%d")
    path = base / y / f"{season}({mname})" / week / inst / day / event
    path.mkdir(parents=True, exist_ok=True)
    logging.info(f"Save directory: {path}")
    return path

def get_view_directory_from_path(selected_file_path: Path) -> Path:
    return selected_file_path.parent


# ── Telegram Queue Worker ─────────────────────────────────────────────────────
def telegram_worker():
    logging.info("Telegram worker thread started.")
    while True:
        try:
            item = telegram_queue.get()
            item_type = item.get('type')

            current_token = TELEGRAM_BOT_TOKEN
            current_chat_id = item.get('chat_id')
            
            if not current_token or current_token == "YOUR_TELEGRAM_BOT_TOKEN_HERE":
                logging.error("Telegram Worker: Bot Token is not configured. Cannot send any messages/photos.")
                telegram_queue.task_done()
                time.sleep(1)
                continue

            if not current_chat_id:
                logging.warning("Telegram Worker: Chat ID is not configured. Cannot send any messages/photos. Please configure it in the GUI.")
                telegram_queue.task_done()
                time.sleep(1)
                continue

            max_retries = 3
            retry_delay_seconds = 5

            for attempt in range(max_retries):
                success = False
                try:
                    if item_type == 'message':
                        message = item['message']
                        url = f"https://api.telegram.org/bot{current_token}/sendMessage"
                        data = {'chat_id': current_chat_id, 'text': message, 'parse_mode': 'Markdown'}
                        logging.debug(f"Worker: Sending text message (Attempt {attempt + 1}/{max_retries}): '{message[:50]}...'")
                        response = requests.post(url, data=data, timeout=10)
                    elif item_type == 'photo':
                        image_path = item['image_path']
                        caption = item['caption']
                        url = f"https://api.telegram.org/bot{current_token}/sendPhoto"
                        with open(image_path, 'rb') as photo_file:
                            files = {'photo': photo_file}
                            data = {'chat_id': current_chat_id, 'caption': caption if caption else ""}
                            logging.debug(f"Worker: Sending photo {image_path.name} (Attempt {attempt + 1}/{max_retries}) with caption '{caption}'")
                            response = requests.post(url, files=files, data=data, timeout=30)
                    else:
                        logging.error(f"Worker: Unknown item type in queue: {item_type}")
                        break

                    logging.debug(f"Worker: Telegram API Response Status Code: {response.status_code}")
                    logging.debug(f"Worker: Telegram API Full Response: {response.text}")
                    
                    response.raise_for_status()
                    result = response.json()
                    if result.get("ok"):
                        logging.info(f"Worker: {item_type.capitalize()} sent successfully to Telegram on attempt 1.")
                        success = True
                        break
                    else:
                        error_desc = result.get('description', 'Unknown API error')
                        logging.error(f"Worker: Failed to send {item_type} (attempt {attempt + 1}): {error_desc}")
                        
                        if "Bad Request: chat not found" in error_desc or \
                           "Unauthorized" in error_desc or \
                           "bot was blocked by the user" in error_desc:
                            logging.critical("Worker: Fatal Telegram API error. Check Bot Token/Chat ID. Exiting worker for this error.")
                            root.after(0, lambda: messagebox.showerror("Telegram Error", f"Fatal Telegram API Error: {error_desc}. Please check your Bot Token and Chat ID settings carefully."))
                            sys.exit(1)
                        
                        if attempt < max_retries - 1:
                            logging.info(f"Worker: Retrying {item_type} in {retry_delay_seconds} seconds due to API error...")
                            time.sleep(retry_delay_seconds)
                        else:
                            logging.error(f"Worker: Max retries ({max_retries}) reached for {item_type} with API error: {error_desc}. Giving up.")
                            
                except requests.exceptions.Timeout:
                    logging.error(f"Worker: Telegram API request timed out for {item_type} (Attempt {attempt + 1}/{max_retries}).")
                    if attempt < max_retries - 1:
                        logging.info(f"Worker: Retrying {item_type} in {retry_delay_seconds} seconds due to timeout...")
                        time.sleep(retry_delay_seconds)
                    else:
                        logging.error(f"Worker: Timeout for {item_type} after {max_retries} attempts. Giving up.")
                except requests.exceptions.ConnectionError as e:
                    logging.error(f"Worker: Network connection error for {item_type} (Attempt {attempt + 1}/{max_retries}): {e}.")
                    if attempt < max_retries - 1:
                        logging.info(f"Worker: Retrying {item_type} in {retry_delay_seconds} seconds due to connection error...")
                        time.sleep(retry_delay_seconds)
                    else:
                        logging.error(f"Worker: Connection error for {item_type} after {max_retries} attempts. Giving up.")
                except requests.exceptions.RequestException as e:
                    logging.error(f"Worker: General Request error sending {item_type} (Attempt {attempt + 1}/{max_retries}): {e}")
                    if attempt < max_retries - 1:
                        logging.info(f"Worker: Retrying {item_type} in {retry_delay_seconds} seconds due to general request error...")
                        time.sleep(retry_delay_seconds)
                    else:
                        logging.error(f"Worker: General Request error for {item_type} after {max_retries} attempts. Giving up.")
                except Exception as e:
                    logging.error(f"Worker: Unexpected error sending {item_type} (Attempt {attempt + 1}/{max_retries}): {e}. Please report this.")
                    if attempt < max_retries - 1:
                        logging.info(f"Worker: Retrying {item_type} in {retry_delay_seconds} seconds due to unexpected error...")
                        time.sleep(retry_delay_seconds)
                    else:
                        logging.error(f"Worker: Unexpected error for {item_type} after {max_retries} attempts. Giving up.")
            
            telegram_queue.task_done()
        except Exception as e:
            logging.error(f"Worker: Unhandled exception in Telegram worker: {e}")
            telegram_queue.task_done()
        time.sleep(0.1)

# ── Function to add items to Telegram Queue ───────────────────────────────────
def add_to_telegram_queue(chat_id: str, item_type: str, **kwargs):
    if not chat_id:
        logging.warning(f"Attempted to add {item_type} to Telegram queue, but Chat ID is empty. Skipping.")
        if root and hasattr(root, 'tk') and app.enable_telegram_send_var.get():
             root.after(0, lambda: messagebox.showwarning("Telegram Setup", "Telegram Chat ID is not configured. Photos/messages will not be sent."))
        return

    item = {'chat_id': chat_id, 'type': item_type}
    item.update(kwargs)
    telegram_queue.put(item)
    logging.info(f"Added {item_type} to Telegram queue. Current queue size: {telegram_queue.qsize()}")

# ── Screenshot & Word Export ─────────────────────────────────────────────────
def take_screenshot_task(event: str, inst: str, mon_names_str: str,
                         telegram_chat_id: str,
                         enable_telegram_send: bool,
                         user_defined_desc: str):
    logging.info(f"Initiating {event} event screenshot capture in background task.")

    try:
        now = datetime.datetime.now()
        save_dir = get_save_directory(event, inst, now)
        ts = now.strftime("%H-%M-%S") # This is the unique timestamp for the set of screenshots
        doc = Document() if event == "Entry" else None
        names = [n.strip() for n in mon_names_str.split(",")]

        event_phrase = "Order Entered" if event == "Entry" else "Order Exited"

        dynamic_description_parts = []
        if user_defined_desc:
            dynamic_description_parts.append(user_defined_desc)
        dynamic_description_parts.append(f"*{event_phrase}* - #{inst} - {now.strftime('%Y-%m-%d %H:%M:%S')}")
        
        dynamic_description = "\n\n".join(dynamic_description_parts)

        sct = mss.mss()

        logging.info("Detected Monitors:")
        current_monitors = get_monitors() 
        for i, mon in enumerate(current_monitors):
            logging.info(f"Monitor {i}: x={mon.x}, y={mon.y}, width={mon.width}, height={mon.height}, is_primary={mon.is_primary}")

        captured_images = []

        if doc:
            doc.add_heading(f"{event} Event Report", level=1)
            doc.add_paragraph(dynamic_description) 
            doc.add_paragraph("")
            doc.add_paragraph("--- Start of Event Screenshots ---") 
            doc.add_page_break() 

        for idx, m in enumerate(current_monitors):
            name = names[idx] if idx < len(names) else f"Monitor {idx+1}"
            
            monitor_region = {
                "top": m.y,
                "left": m.x,
                "width": m.width,
                "height": m.height
            }
            logging.info(f"Capturing region for {name}: {monitor_region}")
            shot = sct.grab(monitor_region)
            
            # ── IMPORTANT CHANGE: Ensure consistent filename format for timestamp matching ──
            # Filenames will now be like "Monitor 1_18-49-32.png" (if name is "Monitor 1")
            # or "Chart_18-49-32.png" (if name is "Chart")
            img_filename = f"{name}_{ts}.png"
            img_path = save_dir / img_filename
            # ────────────────────────────────────────────────────────────────────────────────
            
            mss.tools.to_png(shot.rgb, shot.size, output=str(img_path))
            logging.info(f"Screenshot saved to {img_path}")
            
            captured_images.append({'path': img_path, 'name_for_caption': name, 'monitor_idx': idx}) # Added monitor_idx

            if doc:
                doc.add_paragraph(f"--- {name} ---") 
                doc.add_picture(str(img_path), width=Inches(6))
                doc.add_page_break() 
                
        if doc:
            doc.add_paragraph("--- End of Event Screenshots ---") 
            doc_path = save_dir / f"Trading Journal_{ts}.docx"
            doc.save(str(doc_path))
            logging.info(f"Word document saved to {doc_path}")

        if enable_telegram_send:
            logging.info("Adding Telegram tasks to queue...")
            
            if dynamic_description:
                add_to_telegram_queue(telegram_chat_id, 'message', message=dynamic_description)

            for i, img_data in enumerate(captured_images):
                img_path = img_data['path']
                mon_name_for_caption = img_data['name_for_caption']
                caption = f"#{inst} ({mon_name_for_caption}) - {now.strftime('%H:%M:%S')}" 
                
                logging.info(f"Adding {img_path.name} to Telegram queue with caption: '{caption}'")
                add_to_telegram_queue(telegram_chat_id, 'photo', image_path=img_path, caption=caption)
            
            add_to_telegram_queue(telegram_chat_id, 'message', message="--- End of Event Screenshots ---")
            
        else:
            logging.info("Telegram send is disabled. Skipping photo upload.")
        
        # Update last_view_path after successful save operation (for quick access later)
        if save_dir.exists():
            root.after(0, lambda: app.last_view_path_var.set(str(save_dir)))

    except Exception:
        logging.error("Error in take_screenshot_task:\n" + traceback.format_exc())
        root.after(0, lambda: messagebox.showerror("Error", "Failed to take screenshot.\nSee app.log for details."))

# ── Hotkey Setup ──────────────────────────────────────────────────────────────
def setup_hotkeys(inst_var: tk.StringVar, mon_names_var: tk.StringVar,
                  telegram_chat_id_var: tk.StringVar,
                  enable_telegram_send_var: tk.BooleanVar,
                  default_description_text_widget: tk.Text,
                  last_view_path_var: tk.StringVar): 
    keyboard.add_hotkey('ctrl+shift+e',
                        lambda: threading.Thread(target=take_screenshot_task,
                                       args=("Entry", inst_var.get(), mon_names_var.get(),
                                             telegram_chat_id_var.get(), enable_telegram_send_var.get(),
                                             default_description_text_widget.get("1.0", tk.END).strip()),
                                       daemon=True).start())
    keyboard.add_hotkey('ctrl+shift+x',
                        lambda: threading.Thread(target=take_screenshot_task,
                                       args=("Exit", inst_var.get(), mon_names_var.get(),
                                             telegram_chat_id_var.get(), enable_telegram_send_var.get(),
                                             default_description_text_widget.get("1.0", tk.END).strip()),
                                       daemon=True).start())
    logging.info("Hotkeys bound: Ctrl+Shift+E (Entry), Ctrl+Shift+X (Exit).")

# Event to signal threads to close
close_display_event = threading.Event()

def close_all_image_windows():
    """Signals all active OpenCV display threads to close and then destroys ALL OpenCV windows."""
    logging.info("Signaling all display threads to close...")
    close_display_event.set() # Set the event to signal threads
    
    # Give threads a moment to react and exit their loops
    time.sleep(0.1) 
    
    logging.info("Calling cv2.destroyAllWindows()...")
    cv2.destroyAllWindows() # Destroy all OpenCV windows opened by any thread
    logging.info("cv2.destroyAllWindows() called.")
    
    # We clear the event only after destroying windows, as a safety measure
    close_display_event.clear() # Clear the event for next time


def display_image_fullscreen_on_monitor(image_path: Path, monitor_info, window_name: str):
    """Displays an image fullscreen on a specific monitor, preserving aspect ratio,
    and handling window positioning for multi-monitor setups with DPI scaling.
    """
    try:
        logging.info(f"Attempting to display {image_path.name} on monitor at ({monitor_info.x}, {monitor_info.y})...")
        
        img = cv2.imread(str(image_path))
        if img is None:
            logging.error(f"Could not load image at {image_path}. Check path and file integrity.")
            try:
                pil_img = Image.open(image_path)
                img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
                logging.info(f"Image {image_path.name} loaded successfully using PIL fallback.")
            except Exception as e:
                logging.error(f"PIL fallback also failed for {image_path.name}: {e}")
                return

        img_height, img_width = img.shape[:2]
        mon_width, mon_height = monitor_info.width, monitor_info.height

        # Calculate scaling factors
        scale_w = mon_width / img_width
        scale_h = mon_height / img_height
        
        # Use the smaller scaling factor to fit the image entirely within the monitor
        scale = min(scale_w, scale_h)

        # Calculate new dimensions for the image
        new_width = int(img_width * scale)
        new_height = int(img_height * scale)

        resized_img = cv2.resize(img, (new_width, new_height), interpolation=cv2.INTER_AREA)

        # Create a black background image of the monitor's dimensions
        background = np.zeros((mon_height, mon_width, 3), dtype=np.uint8)

        # Calculate position to center the resized image on the background
        x_offset = (mon_width - new_width) // 2
        y_offset = (mon_height - new_height) // 2

        # Place the resized image onto the black background
        # Ensure slice indices are non-negative and within bounds
        background[max(0, y_offset):min(mon_height, y_offset+new_height), 
                   max(0, x_offset):min(mon_width, x_offset+new_width)] = resized_img

        cv2.namedWindow(window_name, cv2.WINDOW_NORMAL)
        cv2.resizeWindow(window_name, mon_width, mon_height)
        cv2.moveWindow(window_name, monitor_info.x, monitor_info.y)
        
        cv2.imshow(window_name, background) 

        # This loop waits for a global close signal ONLY.
        # Individual ESC key press is ignored, and checking for X button close is removed
        # to prevent unexpected termination of the main application.
        while not close_display_event.is_set():
            cv2.waitKey(10) # Wait for 10ms

    except Exception:
        logging.error(f"Error displaying image {image_path.name}:\n" + traceback.format_exc())
    finally:
        logging.debug(f"Display thread for {window_name} finished. Window destruction handled globally by cv2.destroyAllWindows().")


## NEW FEATURE: View Screenshots - GUI function to ask for input and trigger display
def view_screenshots_gui_task(monitor_names_str: tk.StringVar, view_button_ref: ttk.Button, last_view_path_var: tk.StringVar):
    # Disable the button to prevent multiple simultaneous calls
    root.after(0, lambda: view_button_ref.config(state=tk.DISABLED))
    
    # IMPORTANT: Close any existing image windows before opening new ones
    # This prevents resource leaks and hanging issues.
    close_all_image_windows() 

    selected_file_path = None
    
    try:
        # Get the last viewed path from config, default to BASE_DIR if not set or invalid
        initial_dir = last_view_path_var.get()
        if not Path(initial_dir).is_dir():
            initial_dir = str(get_base_path())
        
        selected_file_path = filedialog.askopenfilename(
            title="Select a screenshot to view its set",
            initialdir=initial_dir, # Start Browse from the last viewed path
            filetypes=[("PNG files", "*.png")],
            parent=root
        )
        
        if not selected_file_path:
            logging.info("File selection cancelled by user.")
            return # User cancelled file selection

        selected_file_path = Path(selected_file_path)
        view_dir = selected_file_path.parent # The directory containing the selected image

        if not view_dir.exists():
            messagebox.showerror("Error", f"Directory not found for the selected image:\n{view_dir}")
            logging.warning(f"Attempted to view screenshots from non-existent directory based on selected file: {view_dir}")
            return

        # ── MODIFIED: Extract timestamp from the selected file name ──────────
        # Example filename: "Monitor 1_18-49-32.png" or "Chart_18-49-32.png"
        # We need to extract "18-49-32"
        match = re.search(r'(\d{2}-\d{2}-\d{2})\.png$', selected_file_path.name)
        if not match:
            messagebox.showerror("Error", "Could not extract timestamp from the selected file name. Please select a valid screenshot.")
            logging.error(f"Failed to extract timestamp from selected file: {selected_file_path.name}")
            return
        
        target_timestamp = match.group(1)
        logging.info(f"Target timestamp extracted from selected file: {target_timestamp}")

        # Now, filter files in the directory to only include those with this timestamp
        all_png_files_in_dir = sorted(list(view_dir.glob("*.png")))
        image_files = []
        for f in all_png_files_in_dir:
            if target_timestamp in f.name:
                image_files.append(f)
        # ──────────────────────────────────────────────────────────────────────

        if not image_files:
            messagebox.showinfo("No Images Found", f"No screenshots matching the timestamp '{target_timestamp}' found in:\n{view_dir}")
            logging.info(f"No screenshots matching timestamp {target_timestamp} found in {view_dir}.")
            return

        # Update last_view_path_var with the directory of the selected file
        root.after(0, lambda: last_view_path_var.set(str(view_dir)))

        # Match images to monitors based on names, if possible
        current_monitors_info = get_monitors()
        monitor_names_list = [n.strip() for n in monitor_names_str.get().split(',')]
        
        images_to_display = []
        
        for img_file in image_files:
            filename_prefix = img_file.name.split('_')[0] 
            matched_monitor_info = None

            try:
                # Try to find a direct match by filename prefix and monitor name
                if filename_prefix in monitor_names_list:
                    monitor_idx_in_list = monitor_names_list.index(filename_prefix)
                    if monitor_idx_in_list < len(current_monitors_info):
                        matched_monitor_info = current_monitors_info[monitor_idx_in_list]
                        logging.debug(f"Direct match for {img_file.name}: {filename_prefix} -> Monitor {monitor_idx_in_list}")
                
                if matched_monitor_info is None: # Fallback to trying to match by index if names are default
                    # If the filename is like "Monitor 1_...", try to parse the index
                    if filename_prefix.startswith("Monitor "):
                        try:
                            mon_num = int(filename_prefix.split(' ')[1])
                            if mon_num > 0 and (mon_num - 1) < len(current_monitors_info):
                                matched_monitor_info = current_monitors_info[mon_num - 1]
                                logging.debug(f"Index match for {img_file.name}: Monitor {mon_num} -> Monitor {mon_num - 1}")
                        except ValueError:
                            pass # Not a simple Monitor N name
                
                if matched_monitor_info is None:
                    # If still no match, log and will try sequential assignment later
                    logging.warning(f"Could not directly map image {img_file.name} to a specific monitor by name/index. Will try sequential.")

            except Exception as e:
                logging.warning(f"Error during image-to-monitor mapping for {img_file.name}: {e}")
                matched_monitor_info = None

            images_to_display.append({'path': img_file, 'monitor_info': matched_monitor_info})

        # Fill in any missing monitor info sequentially based on the order of images found,
        # preferring primary monitor if more images than monitors.
        current_monitor_index = 0
        primary_monitor = next((m for m in current_monitors_info if m.is_primary), current_monitors_info[0] if current_monitors_info else None)
        
        if not primary_monitor: # Should not happen if get_monitors() returns anything
            messagebox.showerror("Error", "No monitors detected. Cannot display screenshots.")
            logging.error("No monitors detected by screeninfo.")
            return

        for item in images_to_display:
            if item['monitor_info'] is None:
                if current_monitor_index < len(current_monitors_info):
                    item['monitor_info'] = current_monitors_info[current_monitor_index]
                    current_monitor_index += 1
                else:
                    # If we run out of distinct monitors, assign remaining images to the primary
                    item['monitor_info'] = primary_monitor
                    logging.warning(f"Assigned {item['path'].name} to primary monitor as no specific or sequential monitor available.")

        if not images_to_display:
            messagebox.showinfo("No Images", f"No displayable images found in:\n{view_dir}")
            return

        # Display images in separate threads
        display_threads = []
        for i, img_data in enumerate(images_to_display):
            monitor_to_use = img_data['monitor_info']
            if monitor_to_use:
                # ── MODIFIED: More descriptive window title for OpenCV windows ──
                window_name = f"{MAIN_WINDOW_TITLE} - {img_data['path'].stem}" 
                t = threading.Thread(target=display_image_fullscreen_on_monitor, 
                                     args=(img_data['path'], monitor_to_use, window_name), 
                                     daemon=True)
                display_threads.append(t)
                t.start()
            else:
                logging.error(f"No monitor info available for {img_data['path'].name}. Skipping display.")

        logging.info("Launched threads for displaying images. User must close main app to close all image windows.")
        messagebox.showinfo("Screenshots Displayed", 
                             "Screenshots are displayed.\n\n"
                             "**To close ALL screenshot windows at once, please close the main application window.**\n"
                             "*(Pressing 'Esc' or clicking 'X' on individual image windows will NOT close them or may cause unexpected behavior.)*")

    except Exception:
        logging.error(f"Error in view_screenshots_gui_task:\n" + traceback.format_exc())
        root.after(0, lambda: messagebox.showerror("Error", "Failed to view screenshots.\nSee app.log for details."))
    finally:
        # Re-enable the button when the task is finished (or an error occurs)
        root.after(0, lambda: view_button_ref.config(state=tk.NORMAL))


# ── Main GUI ─────────────────────────────────────────────────────────────────
def start_gui():
    global root, app
    cfg = load_cfg()
    inst0 = cfg.get("instrument")
    names0 = cfg.get("monitor_names")
    telegram_chat_id0 = cfg.get("telegram_chat_id")
    enable_telegram_send0 = cfg.get("enable_telegram_send")
    default_description0 = cfg.get("default_description")
    last_view_path0 = cfg.get("last_view_path")

    root = tk.Tk()
    # ── MODIFIED: Set Main Window Title ──────────────────────────────────────
    root.title(MAIN_WINDOW_TITLE) 
    # ─────────────────────────────────────────────────────────────────────────
    root.geometry("500x500") 
    root.resizable(False, True) 

    # ── MODIFIED: Set Main Window Icon ───────────────────────────────────────
    if ICON_PATH.exists():
        try:
            root.iconbitmap(str(ICON_PATH))
            # Optional: Attempt to set icon for OpenCV windows (may not work on all systems)
            # This is a bit of a hack, as OpenCV windows are not Tkinter windows.
            # You might need a more complex solution involving PyWin32 for Windows
            # or platform-specific methods for other OS.
            # cv2.setWindowProperty(window_name, cv2.WND_PROP_ASPECT_RATIO, cv2.WINDOW_FREERATIO) 
            # The above line is just an example of a window property, not an icon.
            # There is no direct cv2 function to set window icon globally or by path easily.
            # We're relying on the system default or a hack.
        except tk.TclError as e:
            logging.error(f"Error setting icon: {e}. Make sure '{ICON_PATH.name}' is a valid .ico file.")
            messagebox.showwarning("Icon Error", f"Could not load icon file: {ICON_PATH.name}\nMake sure it's a valid .ico file.")
    else:
        logging.warning(f"Icon file not found at: {ICON_PATH}")
    # ─────────────────────────────────────────────────────────────────────────

    app = type('AppHolder', (object,), {
        'inst_var': tk.StringVar(value=inst0),
        'mon_names_var': tk.StringVar(value=names0),
        'telegram_chat_id_var': tk.StringVar(value=telegram_chat_id0),
        'enable_telegram_send_var': tk.BooleanVar(value=enable_telegram_send0),
        'default_description_text_widget': None,
        'last_view_path_var': tk.StringVar(value=last_view_path0) # New variable for last viewed path
    })()

    f1 = ttk.Frame(root); f1.pack(fill="x", padx=10, pady=5)
    ttk.Label(f1, text="Instrument:").pack(side="left")
    ttk.Combobox(
        f1, values=INSTRUMENTS, textvariable=app.inst_var,
        state="readonly", width=8
    ).pack(side="left", padx=5)
    
    def edit_monitor_names():
        ans = simpledialog.askstring(
            "Monitor Names",
            "Enter names separated by commas (e.g., Chart, Order Book, Time & Sales, News):\nThese names will be used for filenames and Telegram captions.",
            initialvalue=app.mon_names_var.get(),
            parent=root
        )
        if ans:
            app.mon_names_var.set(ans)
    ttk.Button(f1, text="Monitor Names", command=edit_monitor_names) \
        .pack(side="left", padx=5)

    ## NEW FEATURE: View Screenshots - Button for viewing
    # Pass the button itself to the task so it can disable/enable it
    view_screenshots_button = ttk.Button(f1, text="View Screenshots", 
                                         command=lambda: threading.Thread(target=view_screenshots_gui_task, 
                                                                          args=(app.mon_names_var, view_screenshots_button, app.last_view_path_var), 
                                                                          daemon=True).start())
    view_screenshots_button.pack(side="left", padx=5)


    f_desc = ttk.LabelFrame(root, text="Default Event Description"); f_desc.pack(fill="both", expand=True, padx=10, pady=5)
    ttk.Label(f_desc, text="This text will be added to Telegram/Word. '*Order Entered/Exited*' and Instrument/Timestamp will be added automatically.").pack(padx=5, pady=2, anchor="w")
    
    default_description_text_widget = tk.Text(f_desc, wrap="word", height=1, font=("TkDefaultFont", 10), relief="flat", highlightthickness=0)
    default_description_text_widget.pack(fill="both", expand=True, padx=5, pady=2)
    app.default_description_text_widget = default_description_text_widget
    
    default_description_text_widget.insert("1.0", default_description0)

    def adjust_text_height(event=None):
        content_lines = int(default_description_text_widget.index('end-1c').split('.')[0])
        min_height = 1
        max_height = 10 
        calculated_height = max(min_height, min(max_height, content_lines))
        
        if default_description_text_widget.cget("height") != calculated_height:
            default_description_text_widget.config(height=calculated_height)
        
    default_description_text_widget.bind("<KeyRelease>", adjust_text_height)
    default_description_text_widget.bind("<FocusOut>", adjust_text_height) 
    default_description_text_widget.bind("<Configure>", adjust_text_height) 
    
    root.update_idletasks() 
    adjust_text_height()

    f_telegram = ttk.LabelFrame(root, text="Telegram Settings (Bot Token is hardcoded)"); f_telegram.pack(fill="x", padx=10, pady=5)
    
    ttk.Label(f_telegram, text="Private Channel Chat ID:").grid(row=0, column=0, padx=5, pady=2, sticky="w")
    ttk.Entry(f_telegram, textvariable=app.telegram_chat_id_var, width=40).grid(row=0, column=1, padx=5, pady=2, sticky="ew")
    ttk.Label(f_telegram, text="To get Chat ID:\n1. Add bot to private channel as admin.\n2. Send a message in channel.\n3. Go to https://api.telegram.org/botYOUR_BOT_TOKEN/getUpdates in browser.").grid(row=1, column=0, columnspan=2, padx=5, pady=2, sticky="w")
    
    ttk.Checkbutton(f_telegram, text="Enable Telegram Photo Send",
                    variable=app.enable_telegram_send_var).grid(row=2, column=0, columnspan=2, padx=5, pady=5, sticky="w")
    
    f_telegram.grid_columnconfigure(1, weight=1)

    f4 = ttk.Frame(root); f4.pack(fill="x", padx=10, pady=5)
    ttk.Label(f4, text="Ctrl+Shift+E → Entry (PNGs + Word Doc + Telegram)").pack(anchor="w", padx=5)
    ttk.Label(f4, text="Ctrl+Shift+X → Exit (PNGs + Telegram)").pack(anchor="w", padx=5)

    def on_close():
        cfg["instrument"] = app.inst_var.get()
        cfg["monitor_names"] = app.mon_names_var.get()
        cfg["telegram_chat_id"] = app.telegram_chat_id_var.get()
        cfg["enable_telegram_send"] = app.enable_telegram_send_var.get()
        cfg["default_description"] = app.default_description_text_widget.get("1.0", tk.END).strip()
        cfg["last_view_path"] = app.last_view_path_var.get() # Save last viewed path
        save_cfg(cfg)
        close_all_image_windows() # IMPORTANT: Close any open OpenCV windows before exiting
        root.destroy()

    root.protocol("WM_DELETE_WINDOW", on_close)
    
    threading.Thread(target=setup_hotkeys, args=(app.inst_var, app.mon_names_var, app.telegram_chat_id_var, app.enable_telegram_send_var, app.default_description_text_widget, app.last_view_path_var), daemon=True).start()
    threading.Thread(target=telegram_worker, daemon=True).start()
    
    root.mainloop()

# ── Safe Start & Main Guard ───────────────────────────────────────────────────
def safe_start():
    try:
        start_gui()
    except Exception:
        logging.error("Unhandled exception:\n" + traceback.format_exc())
        try:
            temp_root = tk.Tk()
            temp_root.withdraw()
            messagebox.showerror("Fatal Error", "An unexpected critical error occurred.\nSee app.log for details.")
            temp_root.destroy()
        except Exception:
            pass

if __name__ == "__main__":
    safe_start()